import glob
from typing import List
from tqdm import tqdm
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from lxml import html as lxml_html
from rag.service.text import clean_vietnamese_text

class Loader:
    def load_pdf(self, pdf_file: str):
        docs = PyPDFLoader(pdf_file).load()
        for doc in docs:
            doc.page_content = clean_vietnamese_text(doc.page_content)
        return docs

    def load_docx(self, docx_file: str):
        docx = DocxDocument(docx_file)
        parts = [paragraph.text for paragraph in docx.paragraphs]
        for table in docx.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = clean_vietnamese_text("\n".join(parts))
        return [Document(page_content=text, metadata={"source": docx_file})]

    def load_html(self, html_file: str):
        with open(html_file, "rb") as f:
            raw = f.read()

        try:
            tree = lxml_html.fromstring(raw)
            tree.make_links_absolute("http://localhost")
            texts = tree.xpath(
                "//h1/text() | //h2/text() | //h3/text() | //h4/text() | "
                "//p/text() | //li/text() | //td/text() | //th/text() | //blockquote/text()"
            )
            if not texts:
                texts = tree.xpath("string(//body)")  # type: ignore[arg-type]
                if isinstance(texts, str):
                    texts = [texts]
        except Exception:
            text = raw.decode("utf-8", errors="ignore")
            texts = [text]

        cleaned = clean_vietnamese_text("\n".join(str(t) for t in texts))
        return [Document(page_content=cleaned, metadata={"source": html_file})]

    def load_file(self, file_path: str, filename: str | None = None):
        candidate = (filename or file_path).lower()
        if candidate.endswith(".docx"):
            return self.load_docx(file_path)
        if candidate.endswith(".html") or candidate.endswith(".htm"):
            return self.load_html(file_path)
        return self.load_pdf(file_path)

    def load_dir(self, dir_path: str) -> List:
        source_files = glob.glob(f'{dir_path}/*.pdf') + glob.glob(f'{dir_path}/*.docx') + glob.glob(f'{dir_path}/*.html') + glob.glob(f'{dir_path}/*.htm')
        if not source_files:
            raise ValueError(f'No PDF, DOCX, or HTML files found in {dir_path}')
        all_docs = []

        for source_file in tqdm(source_files, desc="Loading source files"):
            try:
                all_docs.extend(self.load_file(source_file))
            except Exception:
                pass
        return all_docs
