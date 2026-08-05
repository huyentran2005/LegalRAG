import logging
from io import BytesIO
from urllib.parse import urlparse
from urllib.parse import parse_qs
import re
from botocore.exceptions import BotoCoreError, ClientError
import httpx
from fastapi import APIRouter, UploadFile, File, Depends, Form, HTTPException, Query, status
from kombu.exceptions import KombuError
from sqlalchemy.orm import Session
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document as DocxDocument
from lxml import html as lxml_html

from app.api.deps import get_current_user
from app.database import get_db
from app.models.document import Document, DocumentStatus
from app.models.chat_session import ChatSession
from app.services.storage_service import upload_file
from app.workers.tasks import process_uploaded_file


router = APIRouter(prefix="/sources", tags=["sources"])
logger = logging.getLogger(__name__)
URL_RE = re.compile(
    r'https?://[^\s]+?(?=[\s\]\)\},;.!?]|$)'
)
URL_LIKE_RE = re.compile(
    r'[\w\W]{0,8}://[^\s]+?(?=[\s\]\)\},;.!?]|$)'
)
DRIVE_FILE_ID_RE = re.compile(
    r'drive\.google\.com/(?:file/d/|open\?id=|uc\?id=|document/d/)([A-Za-z0-9_-]+)',
    re.IGNORECASE,
)
PDF_CONTENT_TYPES = {"application/pdf"}
HTML_CONTENT_TYPES = {"text/html", "application/xhtml+xml"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm"}

def _validate_session(session_id: int | None, db: Session, user_id: int) -> int | None:
    if session_id is None:
        return None
    session = db.get(ChatSession, session_id)
    if not session or session.user_id != user_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Chat session not found")
    return session_id


def _source_payload(document: Document, checked: bool = True) -> dict:
    return {
        "document_id": document.id,
        "session_id": document.session_id,
        "object_key": document.storage_path,
        "name": document.filename,
        "meta": f"{document.page_count} trang",
        "type": document.file_type,
        "status": document.status.value,
        "checked": checked,
    }


def _upload_response(document: Document, linked_documents: list[Document] | None = None) -> dict:
    payload = _source_payload(document)
    payload["linkedSources"] = [
        _source_payload(linked_document)
        for linked_document in (linked_documents or [])
    ]
    return payload


def _extension_from_filename(filename: str) -> str:
    lowered = filename.lower()
    for extension in SUPPORTED_EXTENSIONS:
        if lowered.endswith(extension):
            return extension
    return ""


def _filename_from_url(url: str, content_type: str) -> str:
    parsed = urlparse(url)
    filename = parsed.path.rsplit("/", 1)[-1] or "linked-document"
    if _extension_from_filename(filename):
        return filename
    if content_type in HTML_CONTENT_TYPES:
        return f"{filename}.html"
    if content_type in DOCX_CONTENT_TYPES:
        return f"{filename}.docx"
    return f"{filename}.pdf"


def _is_supported_source(filename: str, content_type: str) -> bool:
    extension = _extension_from_filename(filename)
    return (
        extension in SUPPORTED_EXTENSIONS
        or content_type in PDF_CONTENT_TYPES
        or content_type in HTML_CONTENT_TYPES
        or content_type in DOCX_CONTENT_TYPES
    )


def _extract_links_from_text(text: str) -> list[str]:
    text = _normalize_ocr_text(text)
    links = []
    for match in URL_RE.finditer(text):
        links.append(_repair_url_candidate(match.group()))

    if not links:
        for match in URL_LIKE_RE.finditer(text):
            repaired = _repair_url_candidate(match.group())
            if repaired:
                links.append(repaired)

    links.extend(_extract_drive_links(text))
    return links


def _normalize_ocr_text(text: str) -> str:
    text = text.translate(str.maketrans({
        "ﬁ": "fi",
        "ﬂ": "fl",
        "…": "...",
    }))
    text = text.replace("hƩps://", "https://")
    text = text.replace("hΤps://", "https://")
    text = text.replace("hтtps://", "https://")
    text = text.replace("hps://", "https://")
    text = text.replace("httрs://", "https://")
    text = re.sub(r"https?\s*:\s*//", "https://", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<=https://)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<=drive\.google\.com)/\s+", "/", text, flags=re.IGNORECASE)
    return text


def _repair_url_candidate(candidate: str) -> str:
    candidate = _normalize_ocr_text(candidate).strip()
    candidate = candidate.rstrip(".,;:!?)]}")
    if "://" not in candidate:
        return candidate

    scheme, rest = candidate.split("://", 1)
    scheme_ascii = re.sub(r"[^a-z]", "", scheme.lower())
    if scheme_ascii in {"http", "https", "hps", "htps", "hpps", "hptps"}:
        return f"https://{rest}"
    if scheme_ascii.startswith("h") and scheme_ascii.endswith("ps"):
        return f"https://{rest}"
    return candidate


def _extract_drive_links(text: str) -> list[str]:
    normalized = _normalize_ocr_text(text)
    links: list[str] = []
    for match in DRIVE_FILE_ID_RE.finditer(normalized):
        file_id = match.group(1)
        links.append(f"https://drive.google.com/file/d/{file_id}/view?usp=sharing")

    if "drive.google.com" in normalized:
        compact = re.sub(r"\s+", "", normalized)
        match = re.search(
            r"drive\.google\.com/(?:file/d/|open\?id=|uc\?id=|document/d/)([A-Za-z0-9_-]+)",
            compact,
            re.IGNORECASE,
        )
        if match:
            file_id = match.group(1)
            links.append(f"https://drive.google.com/file/d/{file_id}/view?usp=sharing")

    return list(dict.fromkeys(links))


def _extract_pdf_links(reader: PdfReader, text: str) -> list[str]:
    links = _extract_links_from_text(text)
    for page in reader.pages:
        annotations = page.get("/Annots") or []
        for annotation in annotations:
            try:
                obj = annotation.get_object()
                action = obj.get("/A") or {}
                uri = action.get("/URI")
                if uri:
                    links.append(str(uri).rstrip(".,;:!?)]}"))
            except Exception:
                continue
    return list(dict.fromkeys(links))


def _extract_docx_links(file_obj) -> tuple[int, list[str]]:
    file_obj.seek(0)
    doc = DocxDocument(file_obj)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    links = _extract_links_from_text(text)
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/hyperlink") and rel.target_ref:
            links.append(str(rel.target_ref).rstrip(".,;:!?)]}"))
    return 1, list(dict.fromkeys(links))


def _extract_html_links(html_text: str) -> list[str]:
    try:
        tree = lxml_html.fromstring(html_text)
    except Exception:
        return _extract_links_from_text(html_text)

    links = _extract_links_from_text(html_text)
    for href in tree.xpath("//a/@href | //@src"):
        href = str(href).strip()
        if href.startswith(("http://", "https://")):
            links.append(href.rstrip(".,;:!?)]}"))
    return list(dict.fromkeys(links))


def _inspect_source(file_obj, filename: str, content_type: str) -> tuple[int, list[str]]:
    extension = _extension_from_filename(filename)
    if content_type in PDF_CONTENT_TYPES or extension == ".pdf":
        try:
            file_obj.seek(0)
            reader = PdfReader(file_obj)
            page_count = len(reader.pages)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return page_count, _extract_pdf_links(reader, text)
        except (PdfReadError, ValueError) as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File PDF không hợp lệ hoặc không đọc được.",
            ) from exc

    if content_type in DOCX_CONTENT_TYPES or extension == ".docx":
        try:
            return _extract_docx_links(file_obj)
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File DOCX không hợp lệ hoặc không đọc được.",
            ) from exc

    if content_type in HTML_CONTENT_TYPES or extension in {".html", ".htm"}:
        try:
            file_obj.seek(0)
            raw = file_obj.read()
            if isinstance(raw, bytes):
                raw = raw.decode("utf-8", errors="ignore")
            return 1, _extract_html_links(str(raw))
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File HTML không hợp lệ hoặc không đọc được.",
            ) from exc

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Hiện tại hệ thống chỉ hỗ trợ upload file PDF, DOCX hoặc HTML.",
    )


def _enqueue_document_processing(document_id: int) -> str | None:
    try:
        task = process_uploaded_file.delay(document_id)
        return task.id
    except KombuError:
        logger.exception("Failed to enqueue source processing task")
        return None


def _create_document_from_file_obj(
    db: Session,
    current_user_id: int,
    session_id: int | None,
    file_obj,
    filename: str,
    content_type: str,
    page_count: int,
) -> Document:
    object_key = upload_file(file_obj, filename)
    document = Document(
        owner_id=current_user_id,
        session_id=session_id,
        filename=filename,
        page_count=page_count,
        file_type=content_type,
        storage_path=object_key,
        status=DocumentStatus.PROCESSING,
    )
    db.add(document)
    db.flush()
    return document


def _download_url_source(url: str) -> tuple[BytesIO, str, str]:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("Unsupported URL scheme")
    if "drive.google.com" in parsed.netloc.lower():
        file_id = None
        path_match = re.search(r"/file/d/([A-Za-z0-9_-]+)", parsed.path)
        if path_match:
            file_id = path_match.group(1)
        else:
            qs = parse_qs(parsed.query)
            file_id = (qs.get("id") or [None])[0]
        if file_id:
            url = f"https://drive.google.com/uc?export=download&id={file_id}"
    response = httpx.get(url, follow_redirects=True, timeout=30.0)
    response.raise_for_status()
    content_type = response.headers.get("content-type", "application/pdf").split(";")[0]
    filename = _filename_from_url(url, content_type)
    return BytesIO(response.content), filename, content_type


def _create_linked_documents(
    db: Session,
    current_user_id: int,
    session_id: int | None,
    links: list[str],
) -> list[Document]:
    documents = []
    for link in links:
        try:
            file_obj, filename, content_type = _download_url_source(link)
            if not _is_supported_source(filename, content_type):
                logger.info("Skipping unsupported linked source url=%s content_type=%s", link, content_type)
                continue
            page_count, _nested_links = _inspect_source(file_obj, filename, content_type)
            document = _create_document_from_file_obj(
                db=db,
                current_user_id=current_user_id,
                session_id=session_id,
                file_obj=file_obj,
                filename=filename,
                content_type=content_type,
                page_count=page_count,
            )
            documents.append(document)
        except (httpx.HTTPError, ValueError, HTTPException, BotoCoreError, ClientError, OSError):
            logger.exception("Failed to import linked source url=%s", link)
    return documents

@router.post("/upload")
async def upload_file_endpoint(
    file: UploadFile | None = File(None),
    url: str | None = Form(None),
    sessionId: int | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    session_id = _validate_session(sessionId, db, current_user.id)

    if not file and not url:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Bạn cần chọn file PDF/DOCX/HTML hoặc nhập link phù hợp.",
        )

    if file and url:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chỉ upload một nguồn mỗi lần: file hoặc link.",
        )

    if url:
        try:
            file_obj, filename, content_type = _download_url_source(url)
        except (httpx.HTTPError, ValueError) as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Không tải được tài liệu từ link này.",
            ) from exc
    else:
        filename = file.filename or "uploaded.pdf"  # type: ignore[union-attr]
        content_type = file.content_type or "application/octet-stream"  # type: ignore[union-attr]
        file_obj = file.file  # type: ignore[union-attr]

    if not _is_supported_source(filename, content_type):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Hiện tại hệ thống chỉ hỗ trợ upload file PDF, DOCX hoặc HTML.",
        )

    page_count, links = _inspect_source(file_obj, filename, content_type)

    try:
        document = _create_document_from_file_obj(
            db=db,
            current_user_id=current_user.id,
            session_id=session_id,
            file_obj=file_obj,
            filename=filename,
            content_type=content_type,
            page_count=page_count,
        )

        linked_documents = _create_linked_documents(
            db=db,
            current_user_id=current_user.id,
            session_id=session_id,
            links=links,
        )
        db.commit()
        db.refresh(document)
        for linked_document in linked_documents:
            db.refresh(linked_document)
    except (BotoCoreError, ClientError, OSError) as exc:
        db.rollback()
        logger.exception("Failed to upload source to object storage")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Không thể lưu file vào storage. Kiểm tra MinIO/S3 rồi thử lại.",
        ) from exc
    except Exception:
        db.rollback()
        raise

    _enqueue_document_processing(document.id)
    for linked_document in linked_documents:
        _enqueue_document_processing(linked_document.id)

    return _upload_response(document, linked_documents)

@router.get("/")
async def get_file(
    sessionId: int | None = Query(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    query = db.query(Document).filter(Document.owner_id == current_user.id)
    if sessionId is not None:
        _validate_session(sessionId, db, current_user.id)
        query = query.filter(Document.session_id == sessionId)

    sources = (query
                .order_by(Document.created_at.desc())
                .all()
    )

    result = []

    if len(sources) == 0:
        return []
    
    for s in sources:
        result.append(_source_payload(s))

    return result
    
